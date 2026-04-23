from __future__ import annotations

from io import BytesIO
import json

from app.schemas.admin import RAGLabDocument


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".jsonl",
    ".csv",
    ".tsv",
    ".log",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".htm",
}


def parse_lab_document(filename: str, content: bytes) -> RAGLabDocument:
    suffix = _normalized_suffix(filename)
    if suffix in TEXT_EXTENSIONS:
        return RAGLabDocument(title=filename, content=_decode_text_bytes(content))
    if suffix == ".docx":
        return RAGLabDocument(title=filename, content=_parse_docx(content))
    if suffix == ".xlsx":
        return RAGLabDocument(title=filename, content=_parse_xlsx(content))
    if suffix == ".pdf":
        return RAGLabDocument(title=filename, content=_parse_pdf(content))
    raise ValueError(f"Unsupported file type: {suffix or filename}")


def _normalized_suffix(filename: str) -> str:
    lower = filename.lower().strip()
    if "." not in lower:
        return ""
    return lower[lower.rfind(".") :]


def _decode_text_bytes(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16"):
        try:
            text = content.decode(encoding)
            return _normalize_text(text)
        except UnicodeDecodeError:
            continue
    return _normalize_text(content.decode("utf-8", errors="ignore"))


def _parse_docx(content: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return _normalize_text("\n".join(parts))


def _parse_xlsx(content: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(filename=BytesIO(content), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                parts.append("\t".join(values))
    return _normalize_text("\n".join(parts))


def _parse_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"# Page {index}\n{text}")
    return _normalize_text("\n\n".join(parts))


def _normalize_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        raise ValueError("Parsed file content is empty.")

    if normalized.startswith("{") or normalized.startswith("["):
        try:
            parsed = json.loads(normalized)
            return json.dumps(parsed, ensure_ascii=False, indent=2)
        except Exception:
            return normalized
    return normalized
