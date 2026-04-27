from __future__ import annotations

import asyncio
from collections import OrderedDict
from io import BytesIO
import math
import re
from uuid import uuid4

import jieba

from app.schemas.admin import (
    RAGLabDocument,
    RAGLabQuestionResult,
    RAGLabRequest,
    RAGLabSessionResponse,
    RAGLabSourcePreview,
    RAGLabVariant,
    RAGLabVariantResult,
)
from app.schemas.chat import ChatMessage, IntentDecision, SourceChunk
from app.services.embedding_service import CrossEncoderProvider, SentenceTransformerProvider
from app.services.generator_service import GenerationRequest, QueuedGenerationService

SIMILAR_CHUNK_THRESHOLD = 0.85
MAX_MERGED_SOURCE_LENGTH = 2400
RERANK_SCORE_THRESHOLD = 0.1


class RAGLabService:
    def __init__(
        self,
        embedder_provider: SentenceTransformerProvider | None,
        reranker_provider: CrossEncoderProvider | None,
        generation_service: QueuedGenerationService,
        max_sessions: int = 12,
    ) -> None:
        self._embedder_provider = embedder_provider
        self._reranker_provider = reranker_provider
        self._generation_service = generation_service
        self._max_sessions = max_sessions
        self._sessions: OrderedDict[str, RAGLabSessionResponse] = OrderedDict()
        self._lock = asyncio.Lock()

    async def run(self, request: RAGLabRequest) -> RAGLabSessionResponse:
        if self._embedder_provider is None:
            raise RuntimeError("Embedding model is not available for RAG lab evaluation.")

        session_id = uuid4().hex
        variant_results: list[RAGLabVariantResult] = []
        for variant in request.variants:
            variant_results.append(
                await self._evaluate_variant(
                    variant=variant,
                    questions=request.questions,
                    documents=request.documents,
                )
            )

        session = RAGLabSessionResponse(
            session_id=session_id,
            document_count=len(request.documents),
            question_count=len(request.questions),
            variants=variant_results,
        )
        async with self._lock:
            self._sessions[session_id] = session
            while len(self._sessions) > self._max_sessions:
                self._sessions.popitem(last=False)
        return session

    async def get(self, session_id: str) -> RAGLabSessionResponse | None:
        async with self._lock:
            return self._sessions.get(session_id)

    async def get_variant(self, session_id: str, variant_id: str) -> RAGLabVariantResult | None:
        session = await self.get(session_id)
        if session is None:
            return None
        for variant in session.variants:
            if variant.variant_id == variant_id:
                return variant
        return None

    async def export_excel(self, session_id: str) -> bytes:
        session = await self.get(session_id)
        if session is None:
            raise KeyError(session_id)

        from openpyxl import Workbook

        workbook = Workbook()
        summary_sheet = workbook.active
        summary_sheet.title = "Summary"
        summary_sheet.append(
            ["Session ID", session.session_id, "Documents", session.document_count, "Questions", session.question_count]
        )
        summary_sheet.append([])
        summary_sheet.append(
            [
                "Variant",
                "Question",
                "Answer Preview",
                "Source Titles",
                "Chunk Size",
                "Overlap Ratio",
                "Retrieval K",
                "Retrieval Mode",
                "BM25 Top K",
                "BM25 Title Boost",
                "RRF K",
                "Rerank K",
                "Temperature",
            ]
        )
        for variant in session.variants:
            for question in variant.questions:
                summary_sheet.append(
                    [
                        variant.name,
                        question.question,
                        question.answer_preview,
                        " | ".join(source.title for source in question.sources),
                        variant.chunk_size,
                        variant.chunk_overlap_ratio,
                        variant.retrieval_k,
                        variant.retrieval_mode,
                        variant.bm25_top_k,
                        variant.bm25_title_boost,
                        variant.rrf_k,
                        variant.rerank_k,
                        variant.temperature,
                    ]
                )

        for variant in session.variants:
            sheet = workbook.create_sheet(title=_safe_sheet_title(variant.name))
            sheet.append(
                ["Question", "Answer Full", "Source Title", "Source Score", "Source Preview", "Source Full Content"]
            )
            for question in variant.questions:
                if not question.sources:
                    sheet.append([question.question, question.answer_full, "", "", "", ""])
                    continue
                for source in question.sources:
                    sheet.append(
                        [
                            question.question,
                            question.answer_full,
                            source.title,
                            source.score,
                            source.preview,
                            source.content_full,
                        ]
                    )

        stream = BytesIO()
        workbook.save(stream)
        return stream.getvalue()

    async def export_word(self, session_id: str) -> bytes:
        session = await self.get(session_id)
        if session is None:
            raise KeyError(session_id)

        from docx import Document

        document = Document()
        document.add_heading("RAG Lab Evaluation Report", level=1)
        document.add_paragraph(f"Session ID: {session.session_id}")
        document.add_paragraph(
            f"Documents: {session.document_count} | Questions: {session.question_count}"
        )

        for variant in session.variants:
            document.add_heading(variant.name, level=2)
            document.add_paragraph(
                f"chunk_size={variant.chunk_size}, overlap_ratio={variant.chunk_overlap_ratio}, "
                f"retrieval_k={variant.retrieval_k}, retrieval_mode={variant.retrieval_mode}, "
                f"bm25_top_k={variant.bm25_top_k}, bm25_title_boost={variant.bm25_title_boost}, "
                f"rrf_k={variant.rrf_k}, rerank_k={variant.rerank_k}, temperature={variant.temperature}"
            )
            for index, question in enumerate(variant.questions, start=1):
                document.add_heading(f"Q{index}. {question.question}", level=3)
                document.add_paragraph(question.answer_full or "[No answer]")
                if not question.sources:
                    document.add_paragraph("No sources returned.")
                    continue
                for source_index, source in enumerate(question.sources, start=1):
                    document.add_paragraph(
                        f"[{source_index}] {source.title} | score={source.score:.4f}",
                        style="List Bullet",
                    )
                    document.add_paragraph(source.content_full)

        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()

    async def _evaluate_variant(
        self,
        *,
        variant: RAGLabVariant,
        questions: list[str],
        documents: list[RAGLabDocument],
    ) -> RAGLabVariantResult:
        chunk_overlap = int(round(variant.chunk_size * variant.chunk_overlap_ratio))
        chunk_records = self._build_chunk_records(
            documents=documents,
            chunk_size=variant.chunk_size,
            chunk_overlap=chunk_overlap,
        )
        embeddings = await self._encode_texts([record["content"] for record in chunk_records])
        for record, embedding in zip(chunk_records, embeddings):
            record["embedding"] = embedding

        question_results: list[RAGLabQuestionResult] = []
        for question in questions:
            sources = await self._retrieve_sources(
                question=question,
                chunk_records=chunk_records,
                retrieval_k=variant.retrieval_k,
                retrieval_mode=variant.retrieval_mode,
                bm25_top_k=variant.bm25_top_k,
                bm25_title_boost=variant.bm25_title_boost,
                rrf_k=variant.rrf_k,
                rerank_k=variant.rerank_k,
            )
            answer = await self._generate_answer(
                question=question,
                sources=sources,
                temperature=variant.temperature,
            )
            question_results.append(
                RAGLabQuestionResult(
                    question=question,
                    answer_preview=_shorten(answer, 160),
                    answer_full=answer,
                    sources=[
                        RAGLabSourcePreview(
                            document_id=source.document_id,
                            title=source.title,
                            preview=_shorten(source.content, 180),
                            content_full=source.content,
                            score=source.score,
                            retrieval_mode=source.metadata.get("retrieval_mode"),
                            fusion_sources=source.metadata.get("fusion_sources"),
                            chunk_id=source.metadata.get("chunk_id"),
                            chunk_ids=source.metadata.get("chunk_ids"),
                            merged_chunk_count=source.metadata.get("merged_chunk_count"),
                            citation_index=source.metadata.get("citation_index"),
                        )
                        for source in sources
                    ],
                )
            )

        return RAGLabVariantResult(
            variant_id=variant.variant_id,
            name=variant.name,
            chunk_size=variant.chunk_size,
            chunk_overlap_ratio=variant.chunk_overlap_ratio,
            retrieval_k=variant.retrieval_k,
            retrieval_mode=variant.retrieval_mode,
            bm25_top_k=variant.bm25_top_k,
            bm25_title_boost=variant.bm25_title_boost,
            rrf_k=variant.rrf_k,
            rerank_k=variant.rerank_k,
            temperature=variant.temperature,
            questions=question_results,
        )

    def _build_chunk_records(
        self,
        *,
        documents: list[RAGLabDocument],
        chunk_size: int,
        chunk_overlap: int,
    ) -> list[dict]:
        chunk_records: list[dict] = []
        for document_index, document in enumerate(documents, start=1):
            chunks = _chunk_text(document.content, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            if not chunks:
                chunks = [document.content.strip()]
            for chunk_index, chunk in enumerate(chunks, start=1):
                chunk_records.append(
                    {
                        "document_id": f"lab-{document_index}",
                        "title": document.title,
                        "content": chunk,
                        "chunk_id": f"lab-{document_index}:{chunk_index}",
                        "title_terms": _tokenize_for_bm25(document.title),
                        "content_terms": _tokenize_for_bm25(chunk),
                    }
                )
        return chunk_records

    async def _encode_texts(self, texts: list[str]) -> list[list[float]]:
        embedder = self._embedder_provider.get_model()

        def _encode():
            vectors = embedder.encode(
                texts,
                batch_size=8,
                show_progress_bar=False,
                normalize_embeddings=True,
                convert_to_numpy=True,
            )
            return vectors.tolist()

        return await asyncio.to_thread(_encode)

    async def _retrieve_sources(
        self,
        *,
        question: str,
        chunk_records: list[dict],
        retrieval_k: int,
        retrieval_mode: str,
        bm25_top_k: int,
        bm25_title_boost: float,
        rrf_k: int,
        rerank_k: int,
    ) -> list[SourceChunk]:
        chunks: list[SourceChunk]
        query_vector = None
        if retrieval_mode in {"dense", "hybrid"}:
            query_vector = (await self._encode_texts([question]))[0]

        if retrieval_mode == "dense":
            chunks = _dense_candidates(
                query_vector=query_vector or [],
                chunk_records=chunk_records,
                limit=retrieval_k,
            )
        elif retrieval_mode == "bm25":
            chunks = _bm25_candidates(
                question=question,
                chunk_records=chunk_records,
                limit=bm25_top_k,
                title_boost=bm25_title_boost,
            )[:retrieval_k]
        else:
            dense_chunks = _dense_candidates(
                query_vector=query_vector or [],
                chunk_records=chunk_records,
                limit=retrieval_k,
            )
            bm25_chunks = _bm25_candidates(
                question=question,
                chunk_records=chunk_records,
                limit=bm25_top_k,
                title_boost=bm25_title_boost,
            )
            chunks = _fuse_chunks_with_rrf(
                dense_chunks=dense_chunks,
                bm25_chunks=bm25_chunks,
                rrf_k=rrf_k,
            )

        if rerank_k > 0 and self._reranker_provider is not None and chunks:
            chunks = chunks[:rerank_k]
            reranker = self._reranker_provider.get_model()
            pairs = [(question, f"{chunk.title}\n{chunk.content}") for chunk in chunks]
            scores = await asyncio.to_thread(reranker.predict, pairs)
            reranked = sorted(
                zip(chunks, scores),
                key=lambda item: float(item[1]),
                reverse=True,
            )
            chunks = [
                SourceChunk(
                    document_id=chunk.document_id,
                    title=chunk.title,
                    content=chunk.content,
                    score=float(score),
                    metadata={
                        **chunk.metadata,
                        "retrieval_score": str(chunk.score),
                        "rerank_score": str(float(score)),
                    },
                )
                for chunk, score in reranked
                if float(score) >= RERANK_SCORE_THRESHOLD
            ]

        target_limit = rerank_k if rerank_k > 0 else retrieval_k
        chunks = _govern_lab_sources(chunks, chunk_records)
        return chunks[: min(target_limit, len(chunks))]

    async def _generate_answer(
        self,
        *,
        question: str,
        sources: list[SourceChunk],
        temperature: float,
    ) -> str:
        answer = await self._generation_service.generate(
            GenerationRequest(
                messages=[ChatMessage(role="user", content=question)],
                intent=IntentDecision(
                    intent="knowledge_qa",
                    need_rag=True,
                    rewrite_query=question,
                    rationale="Transient RAG lab evaluation",
                ),
                sources=sources,
                temperature=temperature,
            )
        )
        return answer


def _chunk_text(text: str, *, chunk_size: int, chunk_overlap: int) -> list[str]:
    normalized = text.strip()
    if not normalized:
        return []

    chunks: list[str] = []
    start = 0
    text_length = len(normalized)
    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= text_length:
            break
        start = max(end - chunk_overlap, start + 1)
    return chunks


def _dot(left: list[float], right: list[float]) -> float:
    return float(sum(a * b for a, b in zip(left, right)))


def _dense_candidates(
    *,
    query_vector: list[float],
    chunk_records: list[dict],
    limit: int,
) -> list[SourceChunk]:
    scored = sorted(
        (
            (
                _dot(query_vector, record["embedding"]),
                record,
            )
            for record in chunk_records
        ),
        key=lambda item: item[0],
        reverse=True,
    )[:limit]
    return [
        SourceChunk(
            document_id=record["document_id"],
            title=record["title"],
            content=record["content"],
            score=float(score),
            metadata={
                "chunk_id": record["chunk_id"],
                "retrieval_mode": "dense",
            },
        )
        for score, record in scored
    ]


def _bm25_candidates(
    *,
    question: str,
    chunk_records: list[dict],
    limit: int,
    title_boost: float,
) -> list[SourceChunk]:
    query_terms = _tokenize_for_bm25(question)
    if not query_terms:
        return []

    document_count = len(chunk_records)
    avg_doc_length = sum(
        len(record["title_terms"]) + len(record["content_terms"]) for record in chunk_records
    ) / max(document_count, 1)
    doc_freq: dict[str, int] = {}
    for term in query_terms:
        doc_freq[term] = sum(
            1
            for record in chunk_records
            if term in record["title_terms"] or term in record["content_terms"]
        )

    scored: list[tuple[float, dict]] = []
    for record in chunk_records:
        title_terms = record["title_terms"]
        content_terms = record["content_terms"]
        doc_length = len(title_terms) + len(content_terms)
        score = 0.0
        for term in query_terms:
            freq = content_terms.count(term) + (title_terms.count(term) * title_boost)
            if freq <= 0:
                continue
            idf = math.log(
                1
                + ((document_count - doc_freq.get(term, 0) + 0.5) / (doc_freq.get(term, 0) + 0.5))
            )
            k1 = 1.5
            b = 0.75
            numerator = freq * (k1 + 1)
            denominator = freq + k1 * (1 - b + b * (doc_length / max(avg_doc_length, 1)))
            score += idf * (numerator / max(denominator, 1e-6))
        if score > 0:
            scored.append((score, record))

    scored.sort(key=lambda item: item[0], reverse=True)
    return [
        SourceChunk(
            document_id=record["document_id"],
            title=record["title"],
            content=record["content"],
            score=float(score),
            metadata={
                "chunk_id": record["chunk_id"],
                "retrieval_mode": "bm25",
            },
        )
        for score, record in scored[:limit]
    ]


def _fuse_chunks_with_rrf(
    *,
    dense_chunks: list[SourceChunk],
    bm25_chunks: list[SourceChunk],
    rrf_k: int,
) -> list[SourceChunk]:
    merged: dict[str, dict[str, object]] = {}

    def accumulate(chunks: list[SourceChunk], source_name: str) -> None:
        for rank, chunk in enumerate(chunks, start=1):
            key = chunk.metadata.get("chunk_id") or f"{chunk.document_id}:{hash(chunk.content)}"
            if key not in merged:
                merged[key] = {
                    "chunk": chunk,
                    "score": 0.0,
                    "sources": [],
                    "dense_score": "",
                    "bm25_score": "",
                }
            bucket = merged[key]
            bucket["score"] = float(bucket["score"]) + (1.0 / (rrf_k + rank))
            sources = list(bucket["sources"])
            if source_name not in sources:
                sources.append(source_name)
            bucket["sources"] = sources
            if source_name == "dense":
                bucket["dense_score"] = str(chunk.score)
            if source_name == "bm25":
                bucket["bm25_score"] = str(chunk.score)

    accumulate(dense_chunks, "dense")
    accumulate(bm25_chunks, "bm25")

    fused: list[SourceChunk] = []
    for bucket in sorted(merged.values(), key=lambda item: float(item["score"]), reverse=True):
        chunk = bucket["chunk"]
        metadata = {
            **chunk.metadata,
            "retrieval_mode": "hybrid",
            "fusion_sources": ",".join(bucket["sources"]),
            "fusion_score": str(bucket["score"]),
        }
        if bucket["dense_score"]:
            metadata["dense_score"] = str(bucket["dense_score"])
        if bucket["bm25_score"]:
            metadata["bm25_score"] = str(bucket["bm25_score"])
        fused.append(
            SourceChunk(
                document_id=chunk.document_id,
                title=chunk.title,
                content=chunk.content,
                score=float(bucket["score"]),
                metadata=metadata,
            )
        )
    return fused


def _tokenize_for_bm25(text: str) -> list[str]:
    stopwords = {"什么", "怎么", "如何", "为何", "吗", "么", "呢", "请问"}
    tokens: list[str] = []
    for token in jieba.cut_for_search(text):
        normalized = token.strip()
        if not normalized:
            continue
        if normalized in stopwords:
            continue
        if re.search(r"[0-9A-Za-z\u4e00-\u9fff]", normalized) is None:
            continue
        if normalized.isascii() and len(normalized) == 1 and not normalized.isdigit():
            continue
        tokens.append(normalized)
    return tokens


def _govern_lab_sources(
    chunks: list[SourceChunk],
    chunk_records: list[dict],
) -> list[SourceChunk]:
    vectors = {
        str(record.get("chunk_id")): record.get("embedding")
        for record in chunk_records
        if record.get("chunk_id") and record.get("embedding") is not None
    }
    deduplicated = _deduplicate_lab_chunks(chunks, vectors)
    merged = _merge_lab_adjacent_chunks(deduplicated)
    return [
        _with_source_metadata(chunk, {"citation_index": str(index)})
        for index, chunk in enumerate(merged, start=1)
    ]


def _deduplicate_lab_chunks(
    chunks: list[SourceChunk],
    vectors: dict[str, list[float]],
) -> list[SourceChunk]:
    kept: list[SourceChunk] = []
    kept_vectors: list[list[float] | None] = []
    for chunk in chunks:
        vector = vectors.get(chunk.metadata.get("chunk_id", ""))
        duplicate_index = None
        for index, existing in enumerate(kept):
            if existing.document_id != chunk.document_id:
                continue
            existing_vector = kept_vectors[index]
            if vector is None or existing_vector is None:
                continue
            if _dot(vector, existing_vector) > SIMILAR_CHUNK_THRESHOLD:
                duplicate_index = index
                break
        if duplicate_index is None:
            kept.append(chunk)
            kept_vectors.append(vector)
        elif chunk.score > kept[duplicate_index].score:
            kept[duplicate_index] = chunk
            kept_vectors[duplicate_index] = vector
    return kept


def _merge_lab_adjacent_chunks(chunks: list[SourceChunk]) -> list[SourceChunk]:
    indexed: list[tuple[int, int, SourceChunk]] = []
    passthrough: list[tuple[int, SourceChunk]] = []
    for rank, chunk in enumerate(chunks):
        chunk_index = _parse_lab_chunk_index(chunk)
        if chunk_index is None:
            passthrough.append((rank, chunk))
        else:
            indexed.append((rank, chunk_index, chunk))

    groups: list[dict[str, object]] = []
    for document_id in {chunk.document_id for _, _, chunk in indexed}:
        document_chunks = sorted(
            (item for item in indexed if item[2].document_id == document_id),
            key=lambda item: item[1],
        )
        current: list[tuple[int, int, SourceChunk]] = []
        for item in document_chunks:
            if not current or item[1] - current[-1][1] <= 1:
                current.append(item)
            else:
                groups.append(_build_lab_merge_group(current))
                current = [item]
        if current:
            groups.append(_build_lab_merge_group(current))

    groups.extend(
        {
            "rank": rank,
            "source": _with_source_metadata(
                chunk,
                {
                    "chunk_ids": chunk.metadata.get("chunk_id", ""),
                    "merged_chunk_count": "1",
                },
            ),
        }
        for rank, chunk in passthrough
    )
    groups.sort(key=lambda item: int(item["rank"]))
    return [group["source"] for group in groups]


def _build_lab_merge_group(items: list[tuple[int, int, SourceChunk]]) -> dict[str, object]:
    rank = min(item[0] for item in items)
    chunks_in_order = [item[2] for item in sorted(items, key=lambda item: item[1])]
    best_chunk = max(chunks_in_order, key=lambda chunk: chunk.score)
    chunk_ids = [chunk.metadata.get("chunk_id", "") for chunk in chunks_in_order]
    indexes = [str(item[1]) for item in sorted(items, key=lambda item: item[1])]
    content = "\n\n".join(chunk.content.strip() for chunk in chunks_in_order if chunk.content.strip())
    if len(content) > MAX_MERGED_SOURCE_LENGTH:
        content = content[: MAX_MERGED_SOURCE_LENGTH - 1].rstrip() + "…"
    metadata = {
        **best_chunk.metadata,
        "chunk_ids": ",".join(chunk_id for chunk_id in chunk_ids if chunk_id),
        "chunk_index_range": f"{indexes[0]}-{indexes[-1]}",
        "merged_chunk_count": str(len(chunks_in_order)),
    }
    if len(chunks_in_order) > 1:
        metadata["source_governance"] = "adjacent_merged"
    return {
        "rank": rank,
        "source": SourceChunk(
            document_id=best_chunk.document_id,
            title=best_chunk.title,
            content=content,
            score=best_chunk.score,
            metadata={key: str(value) for key, value in metadata.items()},
        ),
    }


def _parse_lab_chunk_index(chunk: SourceChunk) -> int | None:
    chunk_id = chunk.metadata.get("chunk_id", "")
    if ":" not in chunk_id:
        return None
    try:
        return int(chunk_id.rsplit(":", 1)[-1])
    except ValueError:
        return None


def _with_source_metadata(chunk: SourceChunk, metadata: dict[str, str]) -> SourceChunk:
    return SourceChunk(
        document_id=chunk.document_id,
        title=chunk.title,
        content=chunk.content,
        score=chunk.score,
        metadata={**chunk.metadata, **{key: str(value) for key, value in metadata.items()}},
    )


def _shorten(text: str, limit: int) -> str:
    compact = " ".join(text.split())
    if len(compact) <= limit:
        return compact
    return compact[: limit - 1] + "…"


def _safe_sheet_title(title: str) -> str:
    invalid_chars = set('\\/*?:[]')
    cleaned = "".join("_" if char in invalid_chars else char for char in title).strip()
    cleaned = cleaned[:31]
    return cleaned or "Variant"
